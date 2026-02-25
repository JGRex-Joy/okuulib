"""
Очистка .docx файлов от упоминаний Bizdin.kg прямо в /data.

Использование:
  Один файл:   python -m src.ingestion.clean_docs data/книга.docx
  Вся папка:   python -m src.ingestion.clean_docs data/
"""

import re
import sys
from pathlib import Path
from docx import Document

PATTERN = re.compile(
    r"(www\.)?bizdin\.kg",
    flags=re.IGNORECASE
)


def clean_text(text: str) -> str:
    text = PATTERN.sub("", text)
    text = re.sub(r" {2,}", " ", text).strip()
    return text


def clean_docx(path: Path) -> None:
    doc = Document(str(path))

    for para in doc.paragraphs:
        for run in para.runs:
            cleaned = clean_text(run.text)
            if cleaned != run.text:
                run.text = cleaned

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        cleaned = clean_text(run.text)
                        if cleaned != run.text:
                            run.text = cleaned

    doc.save(str(path))
    print(f"✅ {path.name}")


def main():
    if len(sys.argv) < 2:
        print("❌ Укажи путь до файла или папки.\n")
        print("  Файл:  python -m src.ingestion.clean_docs data/книга.docx")
        print("  Папка: python -m src.ingestion.clean_docs data/")
        sys.exit(1)

    target = Path(sys.argv[1])

    if target.is_file():
        if target.suffix != ".docx":
            print(f"❌ Файл не является .docx: {target}")
            sys.exit(1)
        clean_docx(target)

    elif target.is_dir():
        files = list(target.glob("*.docx"))
        if not files:
            print(f"❌ В папке {target} нет .docx файлов")
            sys.exit(1)
        print(f"🔍 Найдено файлов: {len(files)}\n")
        for f in files:
            clean_docx(f)
        print(f"\n🎉 Готово! Очищено: {len(files)} файл(ов)")

    else:
        print(f"❌ Путь не существует: {target}")
        sys.exit(1)


if __name__ == "__main__":
    main()